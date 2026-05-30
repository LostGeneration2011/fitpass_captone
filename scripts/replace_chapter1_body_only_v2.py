from pathlib import Path
from datetime import datetime
from docx import Document

TARGET = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS_UPDATED_v2.docx")
backup = TARGET.with_name(f"{TARGET.stem}_BACKUP_BEFORE_CH1_BODYONLY_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
backup.write_bytes(TARGET.read_bytes())


def norm(s: str) -> str:
    return " ".join((s or "").strip().lower().split())


def is_heading1(p) -> bool:
    name = (getattr(p.style, "name", "") or "").lower()
    return name.startswith("heading 1")


def is_ch1_heading(p) -> bool:
    t = norm(p.text)
    return (t.startswith("chương 1") or t.startswith("chuong 1")) and is_heading1(p)


def is_ch2_heading(p) -> bool:
    t = norm(p.text)
    return (t.startswith("chương 2") or t.startswith("chuong 2")) and is_heading1(p)


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


doc = Document(str(TARGET))

start_idx = None
for i, p in enumerate(doc.paragraphs):
    if is_ch1_heading(p):
        start_idx = i
        break

if start_idx is None:
    raise RuntimeError("Body Chapter 1 heading not found")

end_idx = None
for i in range(start_idx + 1, len(doc.paragraphs)):
    if is_ch2_heading(doc.paragraphs[i]):
        end_idx = i
        break

if end_idx is None:
    raise RuntimeError("Body Chapter 2 heading not found")

chapter1_heading = doc.paragraphs[start_idx]
chapter1_heading.text = "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI FITPASS"

for i in range(end_idx - 1, start_idx, -1):
    remove_paragraph(doc.paragraphs[i])

lines = [
("", None),
("1.1. Đặt vấn đề", "Heading 2"),
("Trong bối cảnh chuyển đổi số, nhu cầu quản lý phòng gym theo thời gian thực ngày càng cao. Tuy nhiên, nhiều phòng gym vẫn vận hành bằng file thủ công, tin nhắn rời rạc và quy trình thiếu đồng bộ giữa lễ tân, giáo viên và quản lý. Điều này gây ra một số vấn đề:", "Normal"),
("", None),
("Dữ liệu học viên, lớp học, điểm danh dễ lệch nhau giữa các bộ phận.", "Normal"),
("Khó theo dõi tình trạng lớp học theo thời gian thực.", "Normal"),
("Chấm công và tính lương giáo viên mất thời gian, dễ sai sót.", "Normal"),
("Trải nghiệm học viên chưa liền mạch trên mobile.", "Normal"),
("Quản trị vận hành thiếu dashboard tổng quan để ra quyết định nhanh.", "Normal"),
("Vì vậy, cần một nền tảng tập trung giúp kết nối Admin, Teacher, Student trên cùng một hệ thống, vừa dễ dùng vừa có khả năng mở rộng.", "Normal"),
("", None),
("1.2. Giải pháp", "Heading 2"),
("FitPass được xây dựng như một nền tảng quản lý gym đa thành phần, gồm Backend API, Admin Dashboard và Mobile App, tập trung vào trải nghiệm thực tế khi vận hành.", "Normal"),
("", None),
("Các điểm chính của giải pháp:", "Normal"),
("", None),
("Quản lý lớp học theo vòng đời: tạo lớp, duyệt lớp, vận hành session.", "Normal"),
("Điểm danh QR realtime, cập nhật ngay tại buổi học.", "Normal"),
("Phân quyền rõ ràng: ADMIN, TEACHER, STUDENT.", "Normal"),
("Dashboard quản trị tập trung: user, class, session, giao dịch, payroll.", "Normal"),
("Tự động tính lương giáo viên theo dữ liệu session DONE.", "Normal"),
("Đồng bộ dữ liệu giữa web và mobile qua API và Socket.IO.", "Normal"),
("1.3. Kiến trúc hệ thống và công cụ phát triển", "Heading 2"),
("", None),
("1.3.1. Frontend", "Heading 3"),
("", None),
("Admin Web: Next.js 14, React 18.", "Normal"),
("Mobile App: React Native + Expo SDK 54.", "Normal"),
("Giao diện tối ưu cho thao tác nhanh theo từng vai trò người dùng.", "Normal"),
("API client dùng axios, thống nhất chuẩn response.", "Normal"),
("1.3.2. Backend", "Heading 3"),
("", None),
("Node.js + Express.js cho REST API.", "Normal"),
("Prisma ORM kết nối PostgreSQL.", "Normal"),
("JWT Authentication + RBAC Authorization.", "Normal"),
("Socket.IO cho attendance và trạng thái session realtime.", "Normal"),
("Triển khai production trên Railway.", "Normal"),
("1.3.3. Công cụ hỗ trợ phát triển", "Heading 3"),
("", None),
("Git và GitHub để quản lý source code.", "Normal"),
("Visual Studio Code là IDE chính.", "Normal"),
("Postman để test API.", "Normal"),
("Prisma Studio để kiểm tra dữ liệu nhanh.", "Normal"),
("1.3.4. Thiết bị và kiểm thử", "Heading 3"),
("", None),
("Máy phát triển: Windows/macOS.", "Normal"),
("Thiết bị test mobile: Android và iOS.", "Normal"),
("Kiểm thử bằng script API, Postman và manual UI test.", "Normal"),
("1.4. Yêu cầu chức năng chính", "Heading 2"),
("", None),
("Đăng ký, đăng nhập và xác thực tài khoản.", "Normal"),
("Quản lý hồ sơ người dùng theo vai trò.", "Normal"),
("Teacher tạo lớp ở trạng thái PENDING.", "Normal"),
("Admin duyệt hoặc từ chối lớp học.", "Normal"),
("Teacher quản lý session theo lịch.", "Normal"),
("Student xem lớp khả dụng và đăng ký lớp.", "Normal"),
("Teacher tạo mã QR điểm danh theo session.", "Normal"),
("Student quét QR để check-in.", "Normal"),
("Hệ thống cập nhật attendance realtime.", "Normal"),
("Admin và Teacher theo dõi dữ liệu vận hành và lương.", "Normal"),
("Tính lương giáo viên theo session DONE và hourlyRate.", "Normal"),
("1.5. Yêu cầu phi chức năng", "Heading 2"),
("", None),
("Phản hồi API nhanh cho các thao tác phổ biến.", "Normal"),
("Đồng bộ attendance gần thời gian thực.", "Normal"),
("Bảo mật: JWT, hash mật khẩu, kiểm soát truy cập theo role.", "Normal"),
("Dữ liệu nhất quán ở các luồng quan trọng như enrollment, attendance, payroll.", "Normal"),
("Dễ bảo trì và mở rộng theo module.", "Normal"),
("Hoạt động ổn định trên cả web admin và mobile app.", "Normal"),
("1.6. Yêu cầu giao diện (UI/UX)", "Heading 2"),
("", None),
("Thiết kế trực quan, dễ dùng cho người mới.", "Normal"),
("Luồng thao tác ngắn, ít bước cho nghiệp vụ thường xuyên.", "Normal"),
("Màn hình quản trị thể hiện rõ KPI và trạng thái vận hành.", "Normal"),
("Màn hình mobile ưu tiên thao tác nhanh cho lịch học và check-in QR.", "Normal"),
("Điều hướng rõ ràng theo từng vai trò người dùng.", "Normal"),
("1.7. Yêu cầu phần cứng và phần mềm", "Heading 2"),
("", None),
("1.7.1. Thiết bị người dùng", "Heading 3"),
("", None),
("Điện thoại Android hoặc iOS có kết nối internet ổn định.", "Normal"),
("Trình duyệt web hiện đại cho Admin Dashboard.", "Normal"),
("Khuyến nghị RAM đủ để chạy ứng dụng mượt.", "Normal"),
("1.7.2. Môi trường client", "Heading 3"),
("", None),
("Mobile app: cấp quyền camera để quét QR.", "Normal"),
("Web admin: trình duyệt cập nhật phiên bản mới.", "Normal"),
("Kết nối mạng ổn định để đồng bộ realtime.", "Normal"),
("1.7.3. Phần cứng cho phát triển", "Heading 3"),
("", None),
("Laptop/PC có thể chạy Node.js, Next.js, Expo.", "Normal"),
("Thiết bị thật Android/iOS để test mobile.", "Normal"),
("Khuyến nghị RAM từ 8GB trở lên cho quá trình build và debug.", "Normal"),
("1.7.4. Phần mềm cho phát triển", "Heading 3"),
("", None),
("Node.js 20.x, npm 10.x.", "Normal"),
("TypeScript 5.x.", "Normal"),
("PostgreSQL 15.x.", "Normal"),
("Prisma 5.x.", "Normal"),
("Visual Studio Code.", "Normal"),
("Android Studio và/hoặc Xcode cho test mobile.", "Normal"),
("Postman cho kiểm thử API.", "Normal"),
("1.8. Bố cục báo cáo", "Heading 2"),
("Ngoài phần mở đầu và tài liệu tham khảo, nội dung chính gồm 4 chương:", "Normal"),
("", None),
("Chương 1: Tổng quan đề tài.", "Normal"),
("Chương 2: Phân tích và thiết kế hệ thống.", "Normal"),
("Chương 3: Cài đặt, triển khai và kiểm thử.", "Normal"),
("Chương 4: Kết luận và hướng phát triển.", "Normal"),
]

anchor = doc.paragraphs[start_idx + 1]
for text, style in lines:
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style

doc.save(str(TARGET))
print(f"UPDATED={TARGET}")
print(f"BACKUP={backup}")
print(f"START_IDX={start_idx}")
print(f"END_IDX={end_idx}")
