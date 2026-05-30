from pathlib import Path
from datetime import datetime
from docx import Document

TARGET = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS_UPDATED_v2.docx")

if not TARGET.exists():
    raise FileNotFoundError(f"Target file not found: {TARGET}")

backup = TARGET.with_name(f"{TARGET.stem}_BACKUP_BEFORE_CH1_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
backup.write_bytes(TARGET.read_bytes())


def normalize(s: str) -> str:
    return " ".join(s.strip().lower().split())


def starts_chapter_1(text: str) -> bool:
    t = normalize(text)
    return t.startswith("chương 1") or t.startswith("chuong 1")


def starts_chapter_2(text: str) -> bool:
    t = normalize(text)
    return t.startswith("chương 2") or t.startswith("chuong 2")


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


doc = Document(str(TARGET))

start_idx = None
end_idx = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text or ""
    if start_idx is None and starts_chapter_1(txt):
        start_idx = i
        continue
    if start_idx is not None and starts_chapter_2(txt):
        end_idx = i
        break

if start_idx is None or end_idx is None or end_idx <= start_idx:
    raise RuntimeError("Could not locate Chapter 1/Chapter 2 boundaries in the document")

chapter2_para = doc.paragraphs[end_idx]

# Remove current Chapter 1 block only.
for i in range(end_idx - 1, start_idx - 1, -1):
    remove_paragraph(doc.paragraphs[i])

# Insert new Chapter 1 content before Chapter 2.
content = [
    ("Chương 1. TỔNG QUAN ĐỀ TÀI", "Heading 1"),
    ("", None),
    ("1.1 Đặt vấn đề", "Heading 2"),
    ("Trong bối cảnh chuyển đổi số, nhu cầu quản lý phòng gym theo thời gian thực ngày càng cao. Tuy nhiên, nhiều phòng gym vẫn vận hành bằng file thủ công, tin nhắn rời rạc và quy trình thiếu đồng bộ giữa lễ tân, giáo viên và quản lý.", "Normal"),
    ("Điều này dẫn đến sai lệch dữ liệu, chậm xử lý nghiệp vụ, khó theo dõi hiệu quả vận hành và giảm trải nghiệm học viên.", "Normal"),
    ("Vì vậy, cần một nền tảng tập trung để kết nối Admin, Teacher, Student trên cùng hệ thống, vừa dễ dùng vừa có khả năng mở rộng.", "Normal"),
    ("", None),
    ("1.2 Giải pháp", "Heading 2"),
    ("FitPass được xây dựng như một nền tảng quản lý gym đa thành phần gồm Backend API, Admin Dashboard và Mobile App.", "Normal"),
    ("Các điểm chính của giải pháp:", "Normal"),
    ("- Quản lý lớp học theo vòng đời: tạo lớp, duyệt lớp, vận hành session.", "Normal"),
    ("- Điểm danh QR realtime, cập nhật ngay tại buổi học.", "Normal"),
    ("- Phân quyền rõ ràng: ADMIN, TEACHER, STUDENT.", "Normal"),
    ("- Dashboard quản trị tập trung: user, class, session, giao dịch, payroll.", "Normal"),
    ("- Tự động tính lương giáo viên theo dữ liệu session DONE.", "Normal"),
    ("", None),
    ("1.3 Kiến trúc hệ thống và công cụ phát triển", "Heading 2"),
    ("1.3.1 Frontend", "Heading 3"),
    ("- Admin Web: Next.js 14, React 18.", "Normal"),
    ("- Mobile App: React Native + Expo.", "Normal"),
    ("- Giao diện tối ưu theo từng vai trò người dùng.", "Normal"),
    ("1.3.2 Backend", "Heading 3"),
    ("- Node.js + Express.js cho REST API.", "Normal"),
    ("- Prisma ORM kết nối PostgreSQL.", "Normal"),
    ("- JWT Authentication + RBAC Authorization.", "Normal"),
    ("- Socket.IO cho attendance và session realtime.", "Normal"),
    ("1.3.3 Công cụ hỗ trợ", "Heading 3"),
    ("- Git/GitHub, Visual Studio Code, Postman, Prisma Studio.", "Normal"),
    ("", None),
    ("1.4 Yêu cầu chức năng chính", "Heading 2"),
    ("- Đăng ký/đăng nhập/xác thực tài khoản.", "Normal"),
    ("- Quản lý hồ sơ người dùng theo vai trò.", "Normal"),
    ("- Teacher tạo lớp, Admin duyệt lớp.", "Normal"),
    ("- Teacher quản lý session, tạo QR điểm danh.", "Normal"),
    ("- Student đăng ký lớp và quét QR check-in.", "Normal"),
    ("- Tính lương giáo viên theo session DONE.", "Normal"),
    ("", None),
    ("1.5 Yêu cầu phi chức năng", "Heading 2"),
    ("- Đồng bộ dữ liệu nhanh và ổn định.", "Normal"),
    ("- Bảo mật: JWT, hash mật khẩu, phân quyền theo vai trò.", "Normal"),
    ("- Dễ bảo trì và mở rộng.", "Normal"),
    ("- Hoạt động tốt trên web và mobile.", "Normal"),
    ("", None),
    ("1.6 Yêu cầu giao diện (UI/UX)", "Heading 2"),
    ("- Giao diện trực quan, dễ làm quen.", "Normal"),
    ("- Luồng thao tác ngắn cho nghiệp vụ thường xuyên.", "Normal"),
    ("- Màn hình admin rõ KPI, màn hình mobile ưu tiên thao tác nhanh.", "Normal"),
    ("", None),
    ("1.7 Yêu cầu phần cứng và phần mềm", "Heading 2"),
    ("- Thiết bị người dùng: Android/iOS hoặc trình duyệt hiện đại.", "Normal"),
    ("- Môi trường phát triển: Node.js, PostgreSQL, Next.js, Expo, VS Code.", "Normal"),
    ("- Công cụ kiểm thử: Postman, script API.", "Normal"),
    ("", None),
    ("1.8 Bố cục báo cáo", "Heading 2"),
    ("Báo cáo gồm 4 chương:", "Normal"),
    ("- Chương 1: Tổng quan đề tài.", "Normal"),
    ("- Chương 2: Phân tích và thiết kế hệ thống.", "Normal"),
    ("- Chương 3: Cài đặt, triển khai và kiểm thử.", "Normal"),
    ("- Chương 4: Kết luận và hướng phát triển.", "Normal"),
    ("", None),
]

for text, style in content:
    p = chapter2_para.insert_paragraph_before(text)
    if style:
        p.style = style

# Save target doc with only Chapter 1 changed.
doc.save(str(TARGET))
print(f"Updated: {TARGET}")
print(f"Backup: {backup}")
