from pathlib import Path
from datetime import datetime
from docx import Document

TARGET = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS_UPDATED_v2.docx")
if not TARGET.exists():
    raise FileNotFoundError(TARGET)

backup = TARGET.with_name(f"{TARGET.stem}_BACKUP_BEFORE_CH1_EXACT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
backup.write_bytes(TARGET.read_bytes())


def normalize(s: str) -> str:
    return " ".join(s.strip().lower().split())


def starts_ch1(s: str) -> bool:
    t = normalize(s)
    return t.startswith("chương 1") or t.startswith("chuong 1")


def starts_ch2(s: str) -> bool:
    t = normalize(s)
    return t.startswith("chương 2") or t.startswith("chuong 2")


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


doc = Document(str(TARGET))
start_idx = None
end_idx = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text or ""
    if start_idx is None and starts_ch1(txt):
        start_idx = i
        continue
    if start_idx is not None and starts_ch2(txt):
        end_idx = i
        break

if start_idx is None or end_idx is None or end_idx <= start_idx:
    raise RuntimeError("Cannot find Chapter 1/2 boundaries")

chapter2_para = doc.paragraphs[end_idx]

for i in range(end_idx - 1, start_idx - 1, -1):
    remove_paragraph(doc.paragraphs[i])

lines = [
"CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI FITPASS",
"",
"1.1. Đặt vấn đề",
"Trong bối cảnh chuyển đổi số, nhu cầu quản lý phòng gym theo thời gian thực ngày càng cao. Tuy nhiên, nhiều phòng gym vẫn vận hành bằng file thủ công, tin nhắn rời rạc và quy trình thiếu đồng bộ giữa lễ tân, giáo viên và quản lý. Điều này gây ra một số vấn đề:",
"",
"Dữ liệu học viên, lớp học, điểm danh dễ lệch nhau giữa các bộ phận.",
"Khó theo dõi tình trạng lớp học theo thời gian thực.",
"Chấm công và tính lương giáo viên mất thời gian, dễ sai sót.",
"Trải nghiệm học viên chưa liền mạch trên mobile.",
"Quản trị vận hành thiếu dashboard tổng quan để ra quyết định nhanh.",
"Vì vậy, cần một nền tảng tập trung giúp kết nối Admin, Teacher, Student trên cùng một hệ thống, vừa dễ dùng vừa có khả năng mở rộng.",
"",
"1.2. Giải pháp",
"FitPass được xây dựng như một nền tảng quản lý gym đa thành phần, gồm Backend API, Admin Dashboard và Mobile App, tập trung vào trải nghiệm thực tế khi vận hành.",
"",
"Các điểm chính của giải pháp:",
"",
"Quản lý lớp học theo vòng đời: tạo lớp, duyệt lớp, vận hành session.",
"Điểm danh QR realtime, cập nhật ngay tại buổi học.",
"Phân quyền rõ ràng: ADMIN, TEACHER, STUDENT.",
"Dashboard quản trị tập trung: user, class, session, giao dịch, payroll.",
"Tự động tính lương giáo viên theo dữ liệu session DONE.",
"Đồng bộ dữ liệu giữa web và mobile qua API và Socket.IO.",
"1.3. Kiến trúc hệ thống và công cụ phát triển",
"",
"1.3.1. Frontend",
"",
"Admin Web: Next.js 14, React 18.",
"Mobile App: React Native + Expo SDK 54.",
"Giao diện tối ưu cho thao tác nhanh theo từng vai trò người dùng.",
"API client dùng axios, thống nhất chuẩn response.",
"1.3.2. Backend",
"",
"Node.js + Express.js cho REST API.",
"Prisma ORM kết nối PostgreSQL.",
"JWT Authentication + RBAC Authorization.",
"Socket.IO cho attendance và trạng thái session realtime.",
"Triển khai production trên Railway.",
"1.3.3. Công cụ hỗ trợ phát triển",
"",
"Git và GitHub để quản lý source code.",
"Visual Studio Code là IDE chính.",
"Postman để test API.",
"Prisma Studio để kiểm tra dữ liệu nhanh.",
"1.3.4. Thiết bị và kiểm thử",
"",
"Máy phát triển: Windows/macOS.",
"Thiết bị test mobile: Android và iOS.",
"Kiểm thử bằng script API, Postman và manual UI test.",
"1.4. Yêu cầu chức năng chính",
"",
"Đăng ký, đăng nhập và xác thực tài khoản.",
"Quản lý hồ sơ người dùng theo vai trò.",
"Teacher tạo lớp ở trạng thái PENDING.",
"Admin duyệt hoặc từ chối lớp học.",
"Teacher quản lý session theo lịch.",
"Student xem lớp khả dụng và đăng ký lớp.",
"Teacher tạo mã QR điểm danh theo session.",
"Student quét QR để check-in.",
"Hệ thống cập nhật attendance realtime.",
"Admin và Teacher theo dõi dữ liệu vận hành và lương.",
"Tính lương giáo viên theo session DONE và hourlyRate.",
"1.5. Yêu cầu phi chức năng",
"",
"Phản hồi API nhanh cho các thao tác phổ biến.",
"Đồng bộ attendance gần thời gian thực.",
"Bảo mật: JWT, hash mật khẩu, kiểm soát truy cập theo role.",
"Dữ liệu nhất quán ở các luồng quan trọng như enrollment, attendance, payroll.",
"Dễ bảo trì và mở rộng theo module.",
"Hoạt động ổn định trên cả web admin và mobile app.",
"1.6. Yêu cầu giao diện (UI/UX)",
"",
"Thiết kế trực quan, dễ dùng cho người mới.",
"Luồng thao tác ngắn, ít bước cho nghiệp vụ thường xuyên.",
"Màn hình quản trị thể hiện rõ KPI và trạng thái vận hành.",
"Màn hình mobile ưu tiên thao tác nhanh cho lịch học và check-in QR.",
"Điều hướng rõ ràng theo từng vai trò người dùng.",
"1.7. Yêu cầu phần cứng và phần mềm",
"",
"1.7.1. Thiết bị người dùng",
"",
"Điện thoại Android hoặc iOS có kết nối internet ổn định.",
"Trình duyệt web hiện đại cho Admin Dashboard.",
"Khuyến nghị RAM đủ để chạy ứng dụng mượt.",
"1.7.2. Môi trường client",
"",
"Mobile app: cấp quyền camera để quét QR.",
"Web admin: trình duyệt cập nhật phiên bản mới.",
"Kết nối mạng ổn định để đồng bộ realtime.",
"1.7.3. Phần cứng cho phát triển",
"",
"Laptop/PC có thể chạy Node.js, Next.js, Expo.",
"Thiết bị thật Android/iOS để test mobile.",
"Khuyến nghị RAM từ 8GB trở lên cho quá trình build và debug.",
"1.7.4. Phần mềm cho phát triển",
"",
"Node.js 20.x, npm 10.x.",
"TypeScript 5.x.",
"PostgreSQL 15.x.",
"Prisma 5.x.",
"Visual Studio Code.",
"Android Studio và/hoặc Xcode cho test mobile.",
"Postman cho kiểm thử API.",
"1.8. Bố cục báo cáo",
"Ngoài phần mở đầu và tài liệu tham khảo, nội dung chính gồm 4 chương:",
"",
"Chương 1: Tổng quan đề tài.",
"Chương 2: Phân tích và thiết kế hệ thống.",
"Chương 3: Cài đặt, triển khai và kiểm thử.",
"Chương 4: Kết luận và hướng phát triển.",
]

for line in lines:
    p = chapter2_para.insert_paragraph_before(line)
    if line == "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI FITPASS":
        p.style = "Heading 1"
    elif line.startswith("1."):
        p.style = "Heading 2"


doc.save(str(TARGET))
print(f"UPDATED={TARGET}")
print(f"BACKUP={backup}")
