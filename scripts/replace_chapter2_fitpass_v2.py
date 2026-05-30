from pathlib import Path
from docx import Document

TARGET = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS_UPDATED_v2.docx")


def norm(s: str) -> str:
    return " ".join((s or "").strip().lower().split())


def is_heading1(p) -> bool:
    name = (getattr(p.style, "name", "") or "").lower()
    return name.startswith("heading 1")


def is_ch2_heading(p) -> bool:
    t = norm(p.text)
    return (t.startswith("chương 2") or t.startswith("chuong 2") or t.startswith("chapter 2")) and is_heading1(p)


def is_ch3_heading(p) -> bool:
    t = norm(p.text)
    return (t.startswith("chương 3") or t.startswith("chuong 3") or t.startswith("chapter 3")) and is_heading1(p)


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


doc = Document(str(TARGET))

start_idx = None
for i, p in enumerate(doc.paragraphs):
    if is_ch2_heading(p):
        start_idx = i
        break

if start_idx is None:
    raise RuntimeError("Body Chapter 2 heading not found")

end_idx = None
for i in range(start_idx + 1, len(doc.paragraphs)):
    if is_ch3_heading(doc.paragraphs[i]):
        end_idx = i
        break

if end_idx is None:
    raise RuntimeError("Body Chapter 3 heading not found")

# Update chapter title and remove old chapter 2 content.
ch2_heading = doc.paragraphs[start_idx]
ch2_heading.text = "CHAPTER 2: SYSTEM DESIGN ANALYSIS (FITPASS)"
for i in range(end_idx - 1, start_idx, -1):
    remove_paragraph(doc.paragraphs[i])

anchor = doc.paragraphs[start_idx + 1]

# Section headings and intro
p = anchor.insert_paragraph_before("2.1. System requirements analysis")
p.style = "Heading 2"

p = anchor.insert_paragraph_before("2.1.1. Functionality requirements")
p.style = "Heading 3"

rows = [
    ("1", "Sign in with Email", "Users log in to FitPass using email and password."),
    ("2", "Sign in with Google", "Users log in quickly with Google OAuth account."),
    ("3", "Sign up", "New users create account with role STUDENT by default."),
    ("4", "Verify account", "System validates email/account status before full access."),
    ("5", "Forgot password", "System sends reset link/token when user requests password recovery."),
    ("6", "Change password", "Logged-in users change password securely."),
    ("7", "User profile", "Users view and update profile information."),
    ("8", "Role-based access", "System enforces ADMIN/TEACHER/STUDENT permissions."),
    ("9", "Create class", "Teacher creates new class in PENDING status."),
    ("10", "Approve/Reject class", "Admin reviews class and sets APPROVED or REJECTED."),
    ("11", "Class listing", "Students browse available approved classes."),
    ("12", "Class enrollment", "Students enroll into class if conditions are met."),
    ("13", "Session management", "Teacher creates and updates session schedule."),
    ("14", "Session status flow", "Teacher updates UPCOMING -> ONGOING -> DONE."),
    ("15", "Generate attendance QR", "Teacher generates QR token for current session."),
    ("16", "QR check-in", "Student scans QR; system validates token and enrollment."),
    ("17", "Duplicate check-in prevention", "System blocks duplicate attendance for same user/session."),
    ("18", "Realtime attendance sync", "Attendance list updates in realtime via Socket.IO."),
    ("19", "Attendance history", "Users can view attendance records by session/class."),
    ("20", "Payroll calculation", "System calculates teacher salary from DONE sessions and hourlyRate."),
    ("21", "Salary confirmation", "Admin confirms payment status for salary records."),
    ("22", "Dashboard KPI", "Admin dashboard shows users, classes, sessions, transactions, payroll metrics."),
    ("23", "Transaction tracking", "System records transaction data for operational reporting."),
    ("24", "Notifications", "System pushes key updates such as class approval and session changes."),
]

# Title before table
title_p = anchor.insert_paragraph_before("Table 2.1 Functional requirements description (FitPass)")
title_p.style = "Normal"

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "No."
hdr_cells[1].text = "Function name"
hdr_cells[2].text = "Description"

for row in rows:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# Move table to the correct position (before chapter 3 anchor)
anchor._p.addprevious(table._tbl)

p = anchor.insert_paragraph_before("2.1.2. Non-functional requirements")
p.style = "Heading 3"

p = anchor.insert_paragraph_before("2.1.2.1. System requirements")
p.style = "Heading 4"

bullets = [
    "Performance requirement: Fast response for common operations and near realtime attendance updates.",
    "Data consistency requirement: Enrollment, attendance, and payroll data must stay complete and consistent.",
    "Security requirement: Authentication with JWT and authorization with role-based access control.",
    "Design constraints: Backend architecture follows controller-service-data pattern and UML-based analysis.",
    "Extensibility requirement: System can add or remove modules without changing core architecture.",
]
for b in bullets:
    anchor.insert_paragraph_before(f"- {b}")

p = anchor.insert_paragraph_before("2.1.2.2. Interface requirements")
p.style = "Heading 4"

ui_bullets = [
    "Interface must be clear, professional, and easy for new users.",
    "Responsive display for web admin and mobile app screen sizes.",
    "The system uses UTF-8 and supports Vietnamese/English content display.",
    "Date format in system reports: DD/MM/YYYY.",
    "Numeric format in financial views uses clear thousand separators.",
    "Dashboard screens prioritize KPI visibility and quick navigation.",
]
for b in ui_bullets:
    anchor.insert_paragraph_before(f"- {b}")

p = anchor.insert_paragraph_before("2.1.3. Hardware and software requirements")
p.style = "Heading 3"

for b in [
    "Client devices: Android/iOS phones and modern browsers for admin web.",
    "Development environment: Node.js 20.x, npm 10.x, TypeScript 5.x, PostgreSQL 15.x, Prisma 5.x.",
    "Developer tools: Visual Studio Code, Postman, Android Studio and/or Xcode.",
]:
    anchor.insert_paragraph_before(f"- {b}")

p = anchor.insert_paragraph_before("2.2. Chapter summary")
p.style = "Heading 2"
anchor.insert_paragraph_before("Chapter 2 defines FitPass system requirements with functional and non-functional criteria, serving as the baseline for implementation and testing in Chapter 3.")

doc.save(str(TARGET))
print(f"UPDATED={TARGET}")
print(f"CH2_START={start_idx}")
print(f"CH3_START={end_idx}")
