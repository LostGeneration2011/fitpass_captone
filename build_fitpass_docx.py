from collections import OrderedDict
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

INPUT = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS.txt")
OUTPUT = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def format_run(run, bold=False, italic=False, underline=False, size=13):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def add_toc_field(paragraph):
    def append_field_run(parent, kind, text=None):
        run = OxmlElement('w:r')
        if kind == 'text':
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text or ''
            run.append(t)
        else:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), kind)
            run.append(fld)
        parent.append(run)

    append_field_run(paragraph._p, 'begin')
    run = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-4" \\h \\z \\u'
    run.append(instr)
    paragraph._p.append(run)
    append_field_run(paragraph._p, 'separate')

    display = paragraph.add_run('Right-click to update field.')
    format_run(display, italic=True, size=11)

    append_field_run(paragraph._p, 'end')


def add_heading(doc, text, level):
    heading = doc.add_paragraph(style=f'Heading {level}')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    format_run(run, bold=True, size=16 if level == 1 else 13)
    return heading


def add_paragraph(doc, text='', style=None, bold=False, italic=False, underline=False, align=None, size=13):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        format_run(r, bold=bold, italic=italic, underline=underline, size=size)
    return p


def add_text_line(doc, line):
    stripped = line.rstrip('\n')
    if not stripped.strip():
        doc.add_paragraph('')
        return

    if stripped.startswith('================================================================================'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(stripped)
        format_run(r, size=12)
        return

    if 'MỤC LỤC' in stripped:
        add_paragraph(doc, stripped, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_toc_field(p)
        return

    if stripped.startswith('DANH MỤC CÁC KÝ HIỆU') or stripped.startswith('DANH MỤC CÁC BẢNG') or stripped.startswith('DANH MỤC CÁC HÌNH VẼ'):
        add_paragraph(doc, stripped, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
        return

    if re.match(r'^Chương\s+\d+\.\s+', stripped):
        add_heading(doc, stripped, 1)
        return

    if re.match(r'^\d+\.\d+\.\d+\.\d+\s+', stripped):
        add_heading(doc, stripped, 4)
        return

    if re.match(r'^\d+\.\d+\.\d+\s+', stripped):
        add_heading(doc, stripped, 3)
        return

    if re.match(r'^\d+\.\d+\s+', stripped):
        add_heading(doc, stripped, 2)
        return

    if stripped.startswith('Bảng '):
        add_paragraph(doc, stripped, bold=True, size=13)
        return

    p = doc.add_paragraph()
    r = p.add_run(stripped)
    format_run(r, size=13)


def parse_ascii_table(lines, start_idx):
    block = []
    i = start_idx
    while i < len(lines):
        line = lines[i].rstrip('\n')
        block.append(line)
        if line.strip().startswith('└'):
            return block, i + 1
        i += 1
    return block, i


def build_table_from_ascii(doc, caption, block):
    # Convert an ASCII table into a real Word table.
    rows = []
    current = []
    for line in block:
        s = line.strip('\n')
        if not s or s.startswith('┌') or s.startswith('└'):
            continue
        if s.startswith('├') or s.startswith('┼'):
            if current:
                rows.append(current)
                current = []
            continue
        if s.startswith('│'):
            current.append(s)
    if current:
        rows.append(current)

    parsed_rows = []
    max_cols = 0
    for group in rows:
        cells_per_line = []
        for line in group:
            parts = [p.strip() for p in line.split('│')[1:-1]]
            cells_per_line.append(parts)
            max_cols = max(max_cols, len(parts))
        merged = ['' for _ in range(max_cols)]
        for parts in cells_per_line:
            for idx in range(max_cols):
                if idx < len(parts) and parts[idx]:
                    merged[idx] = merged[idx] + ('\n' if merged[idx] else '') + parts[idx]
        parsed_rows.append(merged)

    if not parsed_rows or max_cols == 0:
        return

    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    table.autofit = True
    for r_idx, row in enumerate(parsed_rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ''
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    format_run(run, size=11)
            if r_idx == 0:
                set_cell_shading(cell, 'D9EAF7')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph('')


def parse_usecase_detail(lines, start_idx):
    title = lines[start_idx].strip()
    m = re.match(r'^Bảng\s+(2\.\d+)\s+-\s+(UC-\d+)\s*:\s*(.+)$', title)
    if not m:
        return None, start_idx + 1

    code = m.group(2)
    name = m.group(3)
    data = OrderedDict([
        ('Mã Use Case', code),
        ('Tên Use Case', name),
        ('Actor', ''),
        ('Tiền điều kiện', ''),
        ('Hậu điều kiện', ''),
        ('Luồng chính', ''),
        ('Luồng thay thế', ''),
    ])

    i = start_idx + 1
    current = None
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        if not stripped:
            i += 1
            # stop when blank line followed by next table or next chapter heading
            nxt = ''
            j = i
            while j < len(lines):
                nxt = lines[j].strip()
                if nxt:
                    break
                j += 1
            if not nxt or nxt.startswith('Bảng 2.') or nxt.startswith('2.4.2') or re.match(r'^\d+\.\d+\.\d+\.', nxt) or nxt.startswith('Chương '):
                break
            continue
        if stripped.startswith('Bảng 2.') or stripped.startswith('2.5') or stripped.startswith('2.6') or stripped.startswith('2.7') or stripped.startswith('2.8') or stripped.startswith('2.9') or stripped.startswith('2.10') or stripped.startswith('Chương '):
            break
        if stripped.startswith('Actor:'):
            current = 'Actor'
            data[current] = stripped.split(':', 1)[1].strip()
        elif stripped.startswith('Tiền điều kiện:'):
            current = 'Tiền điều kiện'
            data[current] = stripped.split(':', 1)[1].strip()
        elif stripped.startswith('Hậu điều kiện:'):
            current = 'Hậu điều kiện'
            data[current] = stripped.split(':', 1)[1].strip()
        elif stripped.startswith('Luồng chính:'):
            current = 'Luồng chính'
            data[current] = stripped.split(':', 1)[1].strip()
        elif stripped.startswith('Luồng thay thế:'):
            current = 'Luồng thay thế'
            data[current] = stripped.split(':', 1)[1].strip()
        else:
            if current:
                data[current] += ('\n' if data[current] else '') + stripped
        i += 1

    return data, i


def add_usecase_table(doc, title, data):
    add_paragraph(doc, title, bold=True, size=13)
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    headers = list(data.keys())
    values = list(data.values())
    for i, (h, v) in enumerate(zip(headers, values)):
        table.cell(i, 0).text = h
        table.cell(i, 1).text = v
        set_cell_margins(table.cell(i, 0))
        set_cell_margins(table.cell(i, 1))
        set_cell_shading(table.cell(i, 0), 'F3F6FA')
        for p in table.cell(i, 0).paragraphs + table.cell(i, 1).paragraphs:
            for run in p.runs:
                format_run(run, size=11)
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph('')


def main():
    lines = INPUT.read_text(encoding='utf-8').splitlines()
    doc = Document()

    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    doc.settings.element.append(update_fields)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(13)

    cm = 28.3464567
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2)

    i = 0
    skip_toc = False
    prev_nonempty = ''
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        if stripped == 'MỤC LỤC':
            add_text_line(doc, line)
            skip_toc = True
            prev_nonempty = stripped
            i += 1
            continue

        if skip_toc:
            if stripped.startswith('DANH MỤC CÁC KÝ HIỆU'):
                skip_toc = False
            else:
                i += 1
                continue

        # Use case summary table
        if stripped.startswith('┌') and 'Danh sách Use Case chuẩn hóa' in prev_nonempty:
            block, next_i = parse_ascii_table(lines, i)
            build_table_from_ascii(doc, 'Danh sách Use Case chuẩn hóa trong đồ án', block)
            i = next_i
            prev_nonempty = 'table'
            continue

        # Explicit titled detail tables
        if stripped.startswith('Bảng 2.') and 'UC-' in stripped and ':' in stripped:
            data, next_i = parse_usecase_detail(lines, i)
            if data is not None:
                add_usecase_table(doc, stripped, data)
                i = next_i
                prev_nonempty = stripped
                continue

        # ASCII tables with captions
        if stripped.startswith('┌') and (prev_nonempty.startswith('Bảng 2.18') or prev_nonempty.startswith('Bảng 3.1') or prev_nonempty.startswith('Bảng 3.2') or prev_nonempty.startswith('Bảng 3.3') or prev_nonempty.startswith('Bảng 3.4')):
            block, next_i = parse_ascii_table(lines, i)
            build_table_from_ascii(doc, prev_nonempty, block)
            i = next_i
            prev_nonempty = 'table'
            continue

        if re.match(r'^Chương\s+\d+\.\s+', stripped):
            add_heading(doc, stripped, 1)
            prev_nonempty = stripped
            i += 1
            continue

        if re.match(r'^\d+\.\d+\.\d+\.\d+\s+', stripped):
            add_heading(doc, stripped, 4)
            prev_nonempty = stripped
            i += 1
            continue

        if re.match(r'^\d+\.\d+\.\d+\s+', stripped):
            add_heading(doc, stripped, 3)
            prev_nonempty = stripped
            i += 1
            continue

        if re.match(r'^\d+\.\d+\s+', stripped):
            add_heading(doc, stripped, 2)
            prev_nonempty = stripped
            i += 1
            continue

        # headings / text
        if stripped:
            add_text_line(doc, line)
            prev_nonempty = stripped
        else:
            doc.add_paragraph('')
        i += 1

    # page numbers centered in footer
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('')
    format_run(r)

    temp_path = OUTPUT.with_name('BAO_CAO_DO_AN_FITPASS_build.docx')
    doc.save(str(temp_path))
    print(temp_path)


if __name__ == '__main__':
    main()
