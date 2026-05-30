from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

output_path = r"c:\New folder\fitpass\PHIEU_THEO_DOI_TIEN_DO_FITPASS_12_BUOI.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)

p = doc.add_paragraph("PHIẾU THEO DÕI TIẾN ĐỘ")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

p = doc.add_paragraph("PROJECT")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

doc.add_paragraph("")
doc.add_paragraph("1. Tên đề tài: Hệ thống quản lý phòng gym thông minh FitPass (Web Admin + Mobile App)")
doc.add_paragraph("2. Giảng viên hướng dẫn: ............................................................")
doc.add_paragraph("3. Học viên thực hiện đề tài:")
doc.add_paragraph("   MSHV: ............................      Lớp: ............................")
doc.add_paragraph("   Ngành: Công Nghệ Thông Tin")
doc.add_paragraph("   Chuyên ngành: Phát triển ứng dụng Web và Mobile")

doc.add_paragraph("")

headers = ["Buổi", "Ngày", "Nội dung", "Nhận xét của GVHD (Ký tên)"]
rows = [
    ("1", "27/11/2025", "Khởi động đề tài, xác định phạm vi, mục tiêu, kế hoạch tổng thể; thu thập và phân tích yêu cầu cho Student, Teacher, Admin.", ""),
    ("2", "11/12/2025", "Phân tích nghiệp vụ chi tiết; xây dựng Use Case Diagram, mô tả luồng chính/ngoại lệ; chuẩn hóa danh sách chức năng ưu tiên.", ""),
    ("3", "25/12/2025", "Thiết kế kiến trúc hệ thống: Backend, Admin Dashboard, Mobile App; thiết kế sơ bộ cơ sở dữ liệu và quy ước API.", ""),
    ("4", "08/01/2026", "Thiết kế kỹ thuật chi tiết: ERD, Class Diagram, Sequence Diagram cho đăng nhập, đăng ký lớp, quản lý session, điểm danh QR.", ""),
    ("5", "22/01/2026", "Khởi tạo và triển khai Backend (Node.js/Express + Prisma + PostgreSQL); xây dựng Auth JWT, middleware bảo mật, phân quyền RBAC.", ""),
    ("6", "12/02/2026", "Hoàn thiện API nghiệp vụ chính: user, class, enrollment, session, attendance; kiểm thử API và xử lý lỗi logic nghiệp vụ.", ""),
    ("7", "26/02/2026", "Phát triển Admin Dashboard (Next.js + TypeScript): quản lý người dùng, lớp học, lịch tập, dashboard tổng quan; tích hợp API.", ""),
    ("8", "12/03/2026", "Mở rộng Admin: báo cáo doanh thu, theo dõi attendance, kiểm soát quyền truy cập; tối ưu UX cho quy trình quản trị.", ""),
    ("9", "26/03/2026", "Phát triển Mobile App (Expo React Native): đăng nhập, xem lớp, đăng ký lớp, đặt lịch, theo dõi lịch sử tập luyện.", ""),
    ("10", "16/04/2026", "Tích hợp điểm danh QR và WebSocket real-time; đồng bộ dữ liệu giữa mobile và backend; tối ưu trải nghiệm người dùng.", ""),
    ("11", "07/05/2026", "Hoàn thiện nghiệp vụ nâng cao: payroll giáo viên, thống kê tổng hợp; kiểm thử tích hợp end-to-end và sửa lỗi toàn hệ thống.", ""),
    ("12", "27/05/2026", "UAT, tối ưu hiệu năng, rà soát bảo mật cơ bản; hoàn thiện tài liệu báo cáo, chuẩn bị demo và nghiệm thu đề tài.", ""),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

for row in table.rows:
    row.cells[0].width = Pt(40)
    row.cells[1].width = Pt(90)
    row.cells[2].width = Pt(320)
    row.cells[3].width = Pt(170)

doc.add_paragraph("")
p = doc.add_paragraph("TP. HCM, ngày 27 tháng 05 năm 2026")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph("Giảng viên hướng dẫn")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph("(Ký và ghi rõ họ tên)")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save(output_path)
print(output_path)
