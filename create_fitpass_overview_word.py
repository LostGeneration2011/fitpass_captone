from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = qn(f"w:{edge}")
            element = tc_borders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color="000000", size=8):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "single", "sz": size, "space": 0, "color": color},
                top={"val": "single", "sz": size, "space": 0, "color": color},
                right={"val": "single", "sz": size, "space": 0, "color": color},
                bottom={"val": "single", "sz": size, "space": 0, "color": color},
            )


def clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                top={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
            )


def set_font(run, size=13, bold=False, underline=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def add_text(paragraph, text, size=13, bold=False, underline=False, center=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, underline=underline)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(13)

# Header two-column layout
header = doc.add_table(rows=1, cols=2)
header.alignment = WD_TABLE_ALIGNMENT.CENTER
header.columns[0].width = Cm(8)
header.columns[1].width = Cm(8)
clear_table_borders(header)

left = header.cell(0, 0)
right = header.cell(0, 1)
left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p = left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "VTC ACADEMY - HCMC\n", size=12, bold=True, underline=True)
add_text(p, "KHOA LẬP TRÌNH\n", size=12, bold=True, underline=True)
add_text(p, "____________________", size=12)

p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n", size=12, bold=True, underline=True)
add_text(p, "Độc lập - Tự do - Hạnh phúc\n", size=12, bold=True, underline=True)
add_text(p, "____________________", size=12)

doc.add_paragraph("")
title = doc.add_paragraph()
add_text(title, "TỔNG QUAN DỰ ÁN", size=16, bold=True, underline=True, center=True)
doc.add_paragraph("")

# Main content table
content = doc.add_table(rows=15, cols=2)
content.alignment = WD_TABLE_ALIGNMENT.CENTER
content.columns[0].width = Cm(8)
content.columns[1].width = Cm(8)
set_table_borders(content)

# Merge top rows to single full-width cells
for r in range(0, 12):
    content.cell(r, 0).merge(content.cell(r, 1))

rows_data = [
    "TÊN ĐỀ TÀI: Hệ thống quản lý phòng gym thông minh FitPass",
    "GIẢNG VIÊN HƯỚNG DẪN: ............................................................",
    "THỜI GIAN THỰC HIỆN: 27/11/2025 - 27/05/2026",
    "HỌC VIÊN THỰC HIỆN: Anh Dinh",
    "NỘI DUNG ĐỀ TÀI: Xây dựng nền tảng quản lý phòng gym gồm Backend, Web Admin và Mobile App.",
    "MỤC TIÊU: Số hóa vận hành phòng gym, quản lý lớp học, lịch tập và điểm danh QR theo thời gian thực.",
    "PHẠM VI: Quản lý người dùng (Admin/Teacher/Student), lớp học, session, enrollment, attendance và báo cáo.",
    "PHƯƠNG PHÁP THỰC HIỆN: Node.js + Express + Prisma + PostgreSQL; Next.js cho Admin; Expo React Native cho Mobile.",
    "KẾT QUẢ MONG ĐỢI: Hệ thống hoạt động ổn định, dữ liệu đồng bộ, đáp ứng nghiệp vụ quản trị và trải nghiệm người dùng.",
    "NỘI DUNG LÝ THUYẾT: Kiến trúc 3 lớp, REST API, JWT Authentication, RBAC, ORM, WebSocket realtime.",
    "ỨNG DỤNG THỰC TIỄN: Có thể triển khai cho các phòng gym quy mô nhỏ và vừa, mở rộng theo nhu cầu vận hành.",
    "",
]

for i, text in enumerate(rows_data):
    cell = content.cell(i, 0)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_font(run, size=12, bold=True if i in [0, 1, 2, 3, 4] else False)

# Signature row block (2 columns)
left_sig = content.cell(12, 0)
right_sig = content.cell(12, 1)
left_sig.merge(content.cell(14, 0))
right_sig.merge(content.cell(14, 1))

p = left_sig.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "Xác nhận của giảng viên hướng dẫn\n", size=12, bold=True)
add_text(p, "(Ký và ghi rõ họ tên)", size=12)

p = right_sig.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "TP. Hồ Chí Minh, ngày ..... tháng ..... năm 2026\n", size=12, bold=True)
add_text(p, "Học viên thực hiện\n", size=12, bold=True)
add_text(p, "(Ký và ghi rõ họ tên)\n\n", size=12)
add_text(p, "Anh Dinh", size=12, bold=True)

output_path = r"c:\New folder\fitpass\TONG_QUAN_DU_AN_FITPASS.docx"
doc.save(output_path)
print(output_path)
