from pathlib import Path
from docx import Document

SRC = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS.docx")
OUT = Path(r"c:\New folder\fitpass\BAO_CAO_DO_AN_FITPASS_FIXED.docx")

lines = [
    "2.8.A Bổ sung Activity Diagram",
    "Để hoàn thiện nhóm sơ đồ hành vi trong cùng chương Diagram, nhóm bổ sung Activity Diagram cho ba luồng nghiệp vụ cốt lõi của FITPASS.",
    "AD-01: Activity Diagram - QR Check-in",
    "Luồng xử lý: Teacher bắt đầu session và tạo QR token có thời hạn; Student quét QR trên ứng dụng; Backend xác thực token, kiểm tra điều kiện hợp lệ (không hết hạn, không check-in trùng, session đang diễn ra); nếu hợp lệ thì ghi Attendance và phát sự kiện realtime để cập nhật danh sách điểm danh.",
    "AD-02: Activity Diagram - Class Approval",
    "Luồng xử lý: Teacher gửi yêu cầu tạo lớp ở trạng thái PENDING; Admin nhận danh sách lớp chờ duyệt, xem thông tin và ra quyết định; nếu APPROVED thì lớp được mở cho Student đăng ký, nếu REJECTED thì hệ thống phản hồi kết quả cho Teacher.",
    "AD-03: Activity Diagram - Payroll Calculation",
    "Luồng xử lý: Admin chọn tháng/năm tính lương; hệ thống tổng hợp các session DONE của từng Teacher, tính tổng giờ dạy và số tiền theo hourlyRate; tạo SalaryRecord trạng thái PENDING; sau bước xác nhận thanh toán, trạng thái chuyển sang PAID.",
    "Ý nghĩa bổ sung: Nhóm Activity Diagram giúp mô tả rõ trình tự tác vụ, các điểm rẽ nhánh điều kiện và kết quả đầu ra của từng quy trình, bổ trợ trực tiếp cho bộ Sequence Diagram đã trình bày ở mục 2.8.",
]


def main() -> None:
    doc = Document(str(SRC))

    # Remove manually inserted TOC line (non-clickable) if present.
    for p in doc.paragraphs[:350]:
        t = p.text.strip()
        if t.startswith("2.8.A") and "\t" in p.text:
            p.text = ""

    # Find body insertion block (in Chapter 2 body, not TOC).
    start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if i > 500 and t.startswith("2.8.A"):
            start = i
            break

    if start is None:
        # If somehow missing, insert right before 2.9 in body.
        anchor = None
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if i > 500 and t.startswith("2.9 "):
                anchor = p
                start = i
                break
        if anchor is None:
            raise RuntimeError("Cannot find chapter 2 insertion anchor")
        for text in lines:
            anchor.insert_paragraph_before(text)
    else:
        # Rewrite existing block with clean Vietnamese text.
        for offset, text in enumerate(lines):
            idx = start + offset
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].text = text
            else:
                doc.add_paragraph(text)

    # Ensure section heading style matches 2.8 section heading for TOC inclusion.
    heading_style = None
    for p in doc.paragraphs:
        if p.text.strip() == "2.8 Sequence Diagram":
            heading_style = p.style
            break
    if heading_style is not None and start is not None:
        doc.paragraphs[start].style = heading_style

    # Emphasize AD markers.
    for offset in (2, 4, 6):
        idx = start + offset
        if 0 <= idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            for run in para.runs:
                run.bold = True

    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
